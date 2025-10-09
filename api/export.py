from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Form
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from database.models import get_db, Question, Paper, Student
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import uuid
import zipfile
import io
import glob
import time
import json
from docx import Document
from docx.shared import Inches, Pt

router = APIRouter(prefix="/api", tags=["文件上传和导出"])

# 文件管理配置
MAX_EXPORT_FILES = 20  # 最大保留的导出文件数量
EXPORT_FILE_LIFETIME = 3600  # 导出文件生存时间（秒）

def cleanup_old_exports():
    """清理过期的导出文件"""
    try:
        export_dir = "static/uploads"
        current_time = time.time()
        
        # 获取所有导出文件
        export_files = []
        for ext in ['*.docx', '*.zip']:
            export_files.extend(glob.glob(os.path.join(export_dir, ext)))
        
        # 按修改时间排序，新文件在前
        export_files.sort(key=os.path.getmtime, reverse=True)
        
        files_to_delete = []
        
        # 删除过期的文件
        for file_path in export_files:
            file_age = current_time - os.path.getmtime(file_path)
            if file_age > EXPORT_FILE_LIFETIME:
                files_to_delete.append(file_path)
        
        # 如果文件数量超过限制，删除最旧的文件
        if len(export_files) > MAX_EXPORT_FILES:
            files_to_delete.extend(export_files[MAX_EXPORT_FILES:])
        
        # 执行删除
        for file_path in files_to_delete:
            try:
                os.remove(file_path)
                print(f"已删除过期文件: {file_path}")
            except OSError:
                pass
                
    except Exception as e:
        print(f"清理文件失败: {e}")

# 图片上传
@router.post("/upload")
async def upload_image(
    file: UploadFile = File(...),
    paper_id: Optional[int] = Form(None)
):
    """上传图片
    
    Args:
        file: 上传的图片文件
        paper_id: 试卷ID（可选），如果提供则保存到试卷专属文件夹
    """
    if not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="只允许上传图片文件")
    
    # 生成唯一文件名
    file_extension = os.path.splitext(file.filename)[1]
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    
    # 如果指定了试卷ID，保存到试卷专属文件夹
    if paper_id:
        paper_folder = f"static/uploads/paper_{paper_id}"
        os.makedirs(paper_folder, exist_ok=True)
        file_path = f"{paper_folder}/{unique_filename}"
        url = f"/static/uploads/paper_{paper_id}/{unique_filename}"
    else:
        # 兼容旧版本，保存到根目录
        os.makedirs("static/uploads", exist_ok=True)
        file_path = f"static/uploads/{unique_filename}"
        url = f"/static/uploads/{unique_filename}"
    
    # 保存文件
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    return {"filename": unique_filename, "url": url}

class ExportRequest(BaseModel):
    student_id: int

class ExportMultipleRequest(BaseModel):
    student_ids: List[int]

class ExportSettings(BaseModel):
    # 文档样式
    title_size: int = 16
    question_number_size: int = 12
    text_size: int = 12
    line_spacing: float = 1.5
    
    # 图片设置
    image_scale: float = 1.0
    image_width: float = 5.0  # 英寸
    
    # 间距设置
    question_spacing: int = 0  # 题目后留白行数
    
    # 内容设置
    show_student_info: bool = True
    show_question_text: bool = True
    
    # 单题设置
    question_overrides: Dict[str, dict] = {}

# 获取某题的最终设置（单题优先生效）
def get_question_settings(settings: dict, question_id: int) -> dict:
    qid = str(question_id)
    overrides = settings.get('question_overrides', {})
    result = dict(settings)
    if qid in overrides:
        result.update(overrides[qid])
    return result

@router.post("/export/student/{paper_id}")
async def export_student_word(paper_id: int, request: ExportRequest, db: Session = Depends(get_db)):
    """导出单个学生的错题Word文档"""
    # 清理过期文件
    cleanup_old_exports()
    
    # 获取试卷信息
    paper = db.query(Paper).filter(Paper.id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="试卷不存在")
    
    # 获取学生信息
    student = db.query(Student).filter(Student.id == request.student_id).first()
    if not student:
        raise HTTPException(status_code=404, detail="学生不存在")
    
    # 获取导出设置
    settings = await get_export_settings(paper_id)
    
    # 获取该学生的错题
    questions = db.query(Question).filter(Question.paper_id == paper_id).order_by(Question.question_no).all()
    wrong_questions = []
    
    for question in questions:
        if question.wrong_students:
            wrong_student_ids = [int(id.strip()) for id in question.wrong_students.split(',') if id.strip()]
            if student.id in wrong_student_ids:
                wrong_questions.append(question)
    
    if not wrong_questions:
        raise HTTPException(status_code=404, detail="该学生没有错题")
    
    # 创建Word文档
    doc = Document()
    
    # 应用设置：添加标题
    if settings.get("show_student_info", True):
        title = doc.add_heading(f'{paper.name} - {student.name}({student.student_no}) 错题集', 0)
        title.runs[0].font.size = Pt(settings.get("title_size", 18))
    
    for question in wrong_questions:
        # 获取该题最终设置
        qset = get_question_settings(settings, question.id)
        # 添加题号
        question_heading = doc.add_heading(f'第{question.question_no}题', level=1)
        question_heading.runs[0].font.size = Pt(qset.get("question_number_size", 16))
        # 添加题目文字内容
        if qset.get("show_question_text", True) and question.question_text and question.question_text.strip():
            text_para = doc.add_paragraph(question.question_text)
            text_para.runs[0].font.size = Pt(qset.get("text_size", 14))
        # 添加图片
        if question.image_urls:
            image_urls = question.image_urls.split(',')
            for url in image_urls:
                url = url.strip()
                if url:
                    if url.startswith('/static/uploads/'):
                        image_path = f"static/uploads/{url.split('/')[-1]}"
                    else:
                        image_path = url
                    if os.path.exists(image_path):
                        try:
                            image_width = qset.get("image_width", 5.0) * qset.get("image_scale", 1.0)
                            doc.add_picture(image_path, width=Inches(image_width))
                        except Exception as e:
                            doc.add_paragraph(f"图片加载失败: {url}")
                    else:
                        doc.add_paragraph(f"图片不存在: {url}")
        # 添加题目后留白
        spacing_lines = qset.get("question_spacing", 30)
        for _ in range(spacing_lines):
            doc.add_paragraph()
    
    # 保存到临时文件（格式：班级-学号-姓名.docx）
    temp_filename = f"{student.class_name}-{student.student_no}-{student.name}.docx"
    temp_path = f"static/uploads/{temp_filename}"
    doc.save(temp_path)
    
    return {"download_url": f"/api/download/{temp_filename}"}

@router.get("/download/{filename}")
async def download_file(filename: str):
    """下载文件"""
    file_path = f"static/uploads/{filename}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@router.post("/export/all/{paper_id}")
async def export_all_students(paper_id: int, db: Session = Depends(get_db)):
    """导出所有学生的错题Word文档，按班级分类打包成ZIP"""
    # 清理过期文件
    cleanup_old_exports()
    
    # 获取试卷信息
    paper = db.query(Paper).filter(Paper.id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="试卷不存在")
    
    # 获取导出设置
    settings = await get_export_settings(paper_id)
    
    # 获取所有题目
    questions = db.query(Question).filter(Question.paper_id == paper_id).order_by(Question.question_no).all()
    
    # 获取所有学生，按班级排序
    students = db.query(Student).order_by(Student.class_name, Student.student_no).all()
    
    # 创建ZIP文件
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for student in students:
            # 找到该学生的错题
            wrong_questions = []
            for question in questions:
                if question.wrong_students:
                    wrong_student_ids = [int(id.strip()) for id in question.wrong_students.split(',') if id.strip()]
                    if student.id in wrong_student_ids:
                        wrong_questions.append(question)
            
            if wrong_questions:  # 只为有错题的学生生成文档
                # 创建Word文档
                doc = Document()
                
                # 应用设置：添加标题
                if settings.get("show_student_info", True):
                    title = doc.add_heading(f'{paper.name} - {student.name}({student.student_no}) 错题集', 0)
                    title.runs[0].font.size = Pt(settings.get("title_size", 18))
                
                for question in wrong_questions:
                    qset = get_question_settings(settings, question.id)
                    question_heading = doc.add_heading(f'第{question.question_no}题', level=1)
                    question_heading.runs[0].font.size = Pt(qset.get("question_number_size", 16))
                    if qset.get("show_question_text", True) and question.question_text and question.question_text.strip():
                        text_para = doc.add_paragraph(question.question_text)
                        text_para.runs[0].font.size = Pt(qset.get("text_size", 14))
                    if question.image_urls:
                        image_urls = question.image_urls.split(',')
                        for url in image_urls:
                            url = url.strip()
                            if url:
                                if url.startswith('/static/uploads/'):
                                    image_path = f"static/uploads/{url.split('/')[-1]}"
                                else:
                                    image_path = url
                                if os.path.exists(image_path):
                                    try:
                                        image_width = qset.get("image_width", 5.0) * qset.get("image_scale", 1.0)
                                        doc.add_picture(image_path, width=Inches(image_width))
                                    except Exception as e:
                                        doc.add_paragraph(f"图片加载失败: {url}")
                                else:
                                    doc.add_paragraph(f"图片不存在: {url}")
                    spacing_lines = qset.get("question_spacing", 30)
                    for _ in range(spacing_lines):
                        doc.add_paragraph()
                
                # 保存到内存
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                # 按班级添加到ZIP（使用班级名作为子文件夹，文件名格式：班级-学号-姓名.docx）
                file_path_in_zip = f"{student.class_name}/{student.class_name}-{student.student_no}-{student.name}.docx"
                zip_file.writestr(file_path_in_zip, doc_buffer.getvalue())
    
    zip_buffer.seek(0)
    
    # 保存ZIP文件
    zip_filename = f"{paper.name}_错题集.zip"
    zip_path = f"static/uploads/{zip_filename}"
    
    with open(zip_path, 'wb') as f:
        f.write(zip_buffer.getvalue())
    
    return {"download_url": f"/api/download/{zip_filename}"}

@router.post("/export/students/{paper_id}")
async def export_multiple_students(paper_id: int, request: ExportMultipleRequest, db: Session = Depends(get_db)):
    """导出多个指定学生的错题Word文档，按班级分类打包成ZIP"""
    # 清理过期文件
    cleanup_old_exports()
    
    # 获取试卷信息
    paper = db.query(Paper).filter(Paper.id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="试卷不存在")
    
    # 获取导出设置
    settings = await get_export_settings(paper_id)
    
    # 获取所有题目
    questions = db.query(Question).filter(Question.paper_id == paper_id).order_by(Question.question_no).all()
    
    # 获取指定的学生，按班级排序
    students = db.query(Student).filter(Student.id.in_(request.student_ids)).order_by(Student.class_name, Student.student_no).all()
    
    if not students:
        raise HTTPException(status_code=404, detail="未找到指定的学生")
    
    # 创建ZIP文件
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for student in students:
            # 找到该学生的错题
            wrong_questions = []
            for question in questions:
                if question.wrong_students:
                    wrong_student_ids = [int(id.strip()) for id in question.wrong_students.split(',') if id.strip()]
                    if student.id in wrong_student_ids:
                        wrong_questions.append(question)
            
            if wrong_questions:  # 只为有错题的学生生成文档
                # 创建Word文档
                doc = Document()
                
                # 应用设置：添加标题
                if settings.get("show_student_info", True):
                    title = doc.add_heading(f'{paper.name} - {student.name}({student.student_no}) 错题集', 0)
                    title.runs[0].font.size = Pt(settings.get("title_size", 18))
                
                for question in wrong_questions:
                    qset = get_question_settings(settings, question.id)
                    question_heading = doc.add_heading(f'第{question.question_no}题', level=1)
                    question_heading.runs[0].font.size = Pt(qset.get("question_number_size", 16))
                    if qset.get("show_question_text", True) and question.question_text and question.question_text.strip():
                        text_para = doc.add_paragraph(question.question_text)
                        text_para.runs[0].font.size = Pt(qset.get("text_size", 14))
                    if question.image_urls:
                        image_urls = question.image_urls.split(',')
                        for url in image_urls:
                            url = url.strip()
                            if url:
                                if url.startswith('/static/uploads/'):
                                    image_path = f"static/uploads/{url.split('/')[-1]}"
                                else:
                                    image_path = url
                                if os.path.exists(image_path):
                                    try:
                                        image_width = qset.get("image_width", 5.0) * qset.get("image_scale", 1.0)
                                        doc.add_picture(image_path, width=Inches(image_width))
                                    except Exception as e:
                                        doc.add_paragraph(f"图片加载失败: {url}")
                                else:
                                    doc.add_paragraph(f"图片不存在: {url}")
                    spacing_lines = qset.get("question_spacing", 30)
                    for _ in range(spacing_lines):
                        doc.add_paragraph()
                
                # 保存到内存
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                
                # 按班级添加到ZIP（使用班级名作为子文件夹，文件名格式：班级-学号-姓名.docx）
                file_path_in_zip = f"{student.class_name}/{student.class_name}-{student.student_no}-{student.name}.docx"
                zip_file.writestr(file_path_in_zip, doc_buffer.getvalue())
    
    zip_buffer.seek(0)
    
    # 保存ZIP文件
    zip_filename = f"{paper.name}_错题集_{len(students)}位学生.zip"
    zip_path = f"static/uploads/{zip_filename}"
    
    with open(zip_path, 'wb') as f:
        f.write(zip_buffer.getvalue())
    
    return {"download_url": f"/api/download/{zip_filename}"}

# 导出设置管理接口
@router.post("/export/settings/{paper_id}")
async def save_export_settings(paper_id: int, settings: Dict[str, Any]):
    """保存试卷导出设置（支持完整的设置结构）"""
    settings_dir = "static/export_settings"
    os.makedirs(settings_dir, exist_ok=True)
    
    settings_file = f"{settings_dir}/paper_{paper_id}.json"
    
    with open(settings_file, 'w', encoding='utf-8') as f:
        json.dump(settings, f, ensure_ascii=False, indent=2)
    
    return {"message": "设置保存成功"}

@router.get("/export/preview/{paper_id}")
async def get_paper_preview(paper_id: int, db: Session = Depends(get_db)):
    """获取试卷预览数据（包含所有题目）"""
    # 获取试卷信息
    paper = db.query(Paper).filter(Paper.id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="试卷不存在")
    
    # 获取所有题目
    questions = db.query(Question).filter(Question.paper_id == paper_id).order_by(Question.question_no).all()
    
    # 获取所有学生
    students = db.query(Student).all()
    student_dict = {s.id: {"id": s.id, "name": s.name, "student_no": s.student_no} for s in students}
    
    # 处理题目数据
    questions_data = []
    for question in questions:
        wrong_students = []
        if question.wrong_students:
            wrong_student_ids = [int(id.strip()) for id in question.wrong_students.split(',') if id.strip()]
            wrong_students = [student_dict[sid] for sid in wrong_student_ids if sid in student_dict]
        
        questions_data.append({
            "id": question.id,
            "question_no": question.question_no,
            "question_text": question.question_text or "",
            "image_urls": question.image_urls.split(',') if question.image_urls else [],
            "wrong_students": wrong_students
        })
    
    return {
        "paper": {
            "id": paper.id,
            "name": paper.name
        },
        "questions": questions_data
    }

@router.get("/export/settings/{paper_id}")
async def get_export_settings(paper_id: int):
    """获取试卷导出设置"""
    settings_file = f"static/export_settings/paper_{paper_id}.json"
    
    if os.path.exists(settings_file):
        with open(settings_file, 'r', encoding='utf-8') as f:
            settings = json.load(f)
        # 兼容老配置
        if 'question_overrides' not in settings:
            settings['question_overrides'] = {}
        return settings
    else:
        # 返回默认设置
        return {
            "title_size": 16,
            "question_number_size": 12,
            "text_size": 12,
            "line_spacing": 1.5,
            "image_scale": 1.0,
            "image_width": 5.0,
            "question_spacing": 0,
            "show_student_info": True,
            "show_question_text": True,
            "question_overrides": {}
        }