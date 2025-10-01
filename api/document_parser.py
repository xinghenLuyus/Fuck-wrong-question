from fastapi import APIRouter, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse
import os
import re
import uuid
import tempfile
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
import io
import json
from dataclasses import dataclass

router = APIRouter()

# 支持的文件格式
SUPPORTED_FORMATS = {
    '.pdf': 'PDF文档',
    '.docx': 'Word文档'
}

@dataclass
class ParsedElement:
    """解析元素的数据结构"""
    element_type: str  # 'text' or 'image'
    content: str  # 文本内容或图片路径
    position: tuple  # (x, y, width, height)
    page: int = 0
    order: int = 0
    metadata: Optional[Dict] = None  # 扩展的元数据信息

@dataclass
class ParsedQuestion:
    """解析后的题目数据结构"""
    question_no: int
    question_text: str
    image_urls: List[str]
    raw_content: str = ""
    confidence: float = 1.0

class DocumentParser:
    """文档解析器主类"""
    
    def __init__(self):
        self.upload_dir = "static/uploads"
        os.makedirs(self.upload_dir, exist_ok=True)
    
    def save_temp_file(self, file: UploadFile) -> str:
        """保存上传文件为临时文件"""
        temp_dir = tempfile.gettempdir()
        file_extension = os.path.splitext(file.filename or "temp.tmp")[1]
        temp_filename = f"parse_{uuid.uuid4()}{file_extension}"
        temp_path = os.path.join(temp_dir, temp_filename)
        
        with open(temp_path, "wb") as temp_file:
            content = file.file.read()
            temp_file.write(content)
        
        return temp_path
    
    def cleanup_temp_file(self, file_path: str):
        """清理临时文件"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"清理临时文件失败: {e}")
    
    def save_image(self, image_data: bytes, filename: str) -> str:
        """保存图片到静态目录"""
        image_path = os.path.join(self.upload_dir, filename)
        with open(image_path, "wb") as f:
            f.write(image_data)
        return f"/static/uploads/{filename}"
    
    def extract_word_content_structured(self, docx_path: str) -> List[ParsedElement]:
        """基于文档结构的Word内容提取 - 更精确的方法"""
        doc = Document(docx_path)
        elements = []
        order = 0
        
        # 方法1：按段落样式分组
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
                
            # 获取段落样式信息
            para_style = para.style.name if para.style else 'Normal'
            
            # 检查是否有编号列表
            is_numbered = False
            if para._element.xpath('.//w:numPr'):
                is_numbered = True
            
            # 检查缩进级别
            indent_level = 0
            if para.paragraph_format.left_indent:
                indent_level = int(para.paragraph_format.left_indent.pt // 18)  # 以18pt为一级
            
            # 根据样式和结构判断元素类型
            element_info = {
                'text': text,
                'style': para_style,
                'is_numbered': is_numbered,
                'indent_level': indent_level,
                'para_index': para_idx
            }
            
            elements.append(ParsedElement(
                element_type='text',
                content=text,
                position=(indent_level * 20, para_idx * 20, 100, 20),
                page=0,
                order=order,
                metadata=element_info  # 保存结构信息
            ))
            order += 1
        
        # 提取图片（保持原有逻辑）
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        filename = f"word_img_{uuid.uuid4()}.png"
                        
                        # 转换图片格式
                        try:
                            img = Image.open(io.BytesIO(image_data))
                            if img.mode in ('RGBA', 'LA', 'P'):
                                background = Image.new('RGB', img.size, (255, 255, 255))
                                if img.mode == 'P':
                                    img = img.convert('RGBA')
                                if img.mode == 'RGBA':
                                    background.paste(img, mask=img.split()[-1])
                                else:
                                    background.paste(img)
                                img = background
                            elif img.mode != 'RGB':
                                img = img.convert('RGB')
                            
                            img_buffer = io.BytesIO()
                            img.save(img_buffer, format='PNG', quality=95)
                            final_image_data = img_buffer.getvalue()
                        except Exception:
                            final_image_data = image_data
                        
                        image_url = self.save_image(final_image_data, filename)
                        
                        elements.append(ParsedElement(
                            element_type='image',
                            content=image_url,
                            position=(0, order * 20, 100, 100),
                            page=0,
                            order=order
                        ))
                        order += 1
                        
                    except Exception as e:
                        print(f"处理图片失败: {e}")
                        continue
        except Exception as e:
            print(f"提取图片失败: {e}")
        
        return elements
    
    def structure_based_split(self, elements: List[ParsedElement]) -> List[ParsedQuestion]:
        """基于文档结构的题目切分"""
        text_elements = [e for e in elements if e.element_type == 'text']
        image_elements = [e for e in elements if e.element_type == 'image']
        
        if not text_elements:
            return []
        
        questions = []
        current_question = None
        current_content = []
        
        for element in text_elements:
            text = element.content.strip()
            metadata = getattr(element, 'metadata', {})
            
            # 判断是否是新题目（基于编号列表或缩进）
            is_new_question = False
            
            # 方法1：有编号的列表项
            if metadata.get('is_numbered'):
                is_new_question = True
            
            # 方法2：特定缩进级别且以数字开头
            elif (metadata.get('indent_level', 0) > 0 and 
                  re.match(r'^\d{1,2}[\.．、\)\)]', text)):
                is_new_question = True
            
            # 方法3：显式的题目模式
            elif re.match(r'^\d{1,2}[\.．、\)\)]\s*\S', text):
                is_new_question = True
            
            if is_new_question:
                # 保存上一题
                if current_question is not None and current_content:
                    content = '\n'.join(current_content).strip()
                    if len(content) > 5:
                        questions.append(ParsedQuestion(
                            question_no=current_question,
                            question_text=content,
                            image_urls=[],
                            raw_content=content,
                            confidence=0.95
                        ))
                
                # 开始新题
                current_question = len(questions) + 1
                current_content = [text]
            else:
                # 继续当前题目
                if current_question is not None:
                    current_content.append(text)
        
        # 处理最后一题
        if current_question is not None and current_content:
            content = '\n'.join(current_content).strip()
            if len(content) > 5:
                questions.append(ParsedQuestion(
                    question_no=current_question,
                    question_text=content,
                    image_urls=[],
                    raw_content=content,
                    confidence=0.95
                ))
        
        # 分配图片
        if image_elements and questions:
            images_per_question = len(image_elements) // len(questions)
            current_img_idx = 0
            
            for i, question in enumerate(questions):
                img_count = images_per_question
                if i < len(image_elements) % len(questions):
                    img_count += 1
                
                if img_count > 0:
                    end_idx = min(current_img_idx + img_count, len(image_elements))
                    question.image_urls = [img.content for img in image_elements[current_img_idx:end_idx]]
                    current_img_idx = end_idx
        
        return questions
    
    def visual_based_split(self, file_path: str) -> List[ParsedQuestion]:
        """基于AI视觉的题目切分 - 使用OCR和布局分析"""
        questions = []
        
        try:
            # 方案A: 使用PDF的视觉布局信息
            if file_path.lower().endswith('.pdf'):
                doc = fitz.open(file_path)
                
                for page_num in range(doc.page_count):
                    page = doc[page_num]
                    
                    # 获取所有文本块的位置信息
                    blocks = page.get_text("dict")["blocks"]
                    
                    # 按位置分组（Y轴聚类）
                    text_lines = []
                    for block in blocks:
                        if "lines" in block:
                            for line in block["lines"]:
                                line_text = ""
                                line_bbox = line["bbox"]
                                for span in line["spans"]:
                                    line_text += span["text"]
                                
                                if line_text.strip():
                                    text_lines.append({
                                        'text': line_text.strip(),
                                        'bbox': line_bbox,
                                        'y': line_bbox[1]  # Y坐标用于排序
                                    })
                    
                    # 按Y坐标排序
                    text_lines.sort(key=lambda x: x['y'])
                    
                    # 基于空白间隙切分题目
                    current_question = []
                    prev_y = 0
                    
                    for line in text_lines:
                        y_gap = line['y'] - prev_y
                        
                        # 如果间隙超过阈值，且当前行以数字开头，则开始新题
                        if (y_gap > 10 and  # 10像素间隙
                            re.match(r'^\d{1,2}[\.．、\)\)]', line['text']) and
                            current_question):  # 不是第一行
                            
                            # 保存上一题
                            if current_question:
                                content = '\n'.join([l['text'] for l in current_question])
                                if len(content.strip()) > 5:
                                    questions.append(ParsedQuestion(
                                        question_no=len(questions) + 1,
                                        question_text=content,
                                        image_urls=[],
                                        raw_content=content,
                                        confidence=0.85
                                    ))
                            
                            # 开始新题
                            current_question = [line]
                        else:
                            current_question.append(line)
                        
                        prev_y = line['y']
                    
                    # 处理最后一题
                    if current_question:
                        content = '\n'.join([l['text'] for l in current_question])
                        if len(content.strip()) > 5:
                            questions.append(ParsedQuestion(
                                question_no=len(questions) + 1,
                                question_text=content,
                                image_urls=[],
                                raw_content=content,
                                confidence=0.85
                            ))
                
                doc.close()
            
            # 方案B: Word文档的视觉分析（简化版）
            elif file_path.lower().endswith('.docx'):
                # 这里可以集成更高级的OCR工具，如PaddleOCR或Tesseract
                # 目前回退到结构化分析
                elements = self.extract_word_content_structured(file_path)
                questions = self.structure_based_split(elements)
                
        except Exception as e:
            print(f"视觉分析失败: {e}")
            return []
        
        return questions
    
    def manual_split_assistance(self, elements: List[ParsedElement]) -> Dict:
        """手动辅助切分 - 返回用于手动调整的数据结构"""
        text_elements = [e for e in elements if e.element_type == 'text']
        image_elements = [e for e in elements if e.element_type == 'image']
        
        # 返回结构化数据，供前端手动调整
        return {
            'text_blocks': [
                {
                    'id': i,
                    'content': e.content,
                    'position': e.position,
                    'order': e.order,
                    'suggested_question': self._suggest_question_boundary(e.content)
                }
                for i, e in enumerate(text_elements)
            ],
            'images': [
                {
                    'id': i,
                    'url': e.content,
                    'position': e.position,
                    'order': e.order
                }
                for i, e in enumerate(image_elements)
            ],
            'suggestions': self._generate_split_suggestions(text_elements)
        }
    
    def _suggest_question_boundary(self, text: str) -> bool:
        """建议是否为题目边界"""
        # 简单的启发式判断
        if re.match(r'^\d{1,2}[\.．、\)\)]', text.strip()):
            return True
        if any(keyword in text for keyword in ['？', '?', '填空', '选择', '计算']):
            return True
        return False
    
    def _generate_split_suggestions(self, text_elements: List[ParsedElement]) -> List[str]:
        """生成切分建议"""
        suggestions = []
        
        # 统计分析
        total_text = '\n'.join(e.content for e in text_elements)
        
        # 可能的题目数量
        potential_questions = len(re.findall(r'\d{1,2}[\.．、\)\)]', total_text))
        if potential_questions > 0:
            suggestions.append(f"可能包含 {potential_questions} 道题目")
        
        # 大题数量
        section_headers = len(re.findall(r'[一二三四五六七八九十]+[、．.]', total_text))
        if section_headers > 0:
            suggestions.append(f"可能包含 {section_headers} 个大题部分")
        
        return suggestions
        """从 Word 文档提取内容"""
        doc = Document(docx_path)
        elements = []
        order = 0
        
        # 按文档顺序处理所有元素
        for para_idx, para in enumerate(doc.paragraphs):
            # 提取段落文本
            text = para.text.strip()
            if text:
                elements.append(ParsedElement(
                    element_type='text',
                    content=text,
                    position=(0, para_idx * 20, 100, 20),  # 估算位置
                    page=0,
                    order=order
                ))
                order += 1
        
        # 提取所有图片（从关系中获取）
        try:
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        # 生成唯一文件名
                        image_ext = os.path.splitext(rel.target_ref)[1] or '.png'
                        filename = f"word_img_{uuid.uuid4()}.png"
                        
                        # 转换图片格式为PNG确保兼容性
                        try:
                            img = Image.open(io.BytesIO(image_data))
                            # 确保图片为RGB模式
                            if img.mode in ('RGBA', 'LA', 'P'):
                                background = Image.new('RGB', img.size, (255, 255, 255))
                                if img.mode == 'P':
                                    img = img.convert('RGBA')
                                if img.mode == 'RGBA':
                                    background.paste(img, mask=img.split()[-1])
                                else:
                                    background.paste(img)
                                img = background
                            elif img.mode != 'RGB':
                                img = img.convert('RGB')
                            
                            # 保存为PNG
                            img_buffer = io.BytesIO()
                            img.save(img_buffer, format='PNG', quality=95)
                            final_image_data = img_buffer.getvalue()
                        except Exception as img_error:
                            print(f"图片转换失败，使用原始数据: {img_error}")
                            final_image_data = image_data
                        
                        # 保存图片
                        image_url = self.save_image(final_image_data, filename)
                        
                        elements.append(ParsedElement(
                            element_type='image',
                            content=image_url,
                            position=(0, order * 20, 100, 100),  # 估算位置
                            page=0,
                            order=order
                        ))
                        order += 1
                        
                    except Exception as e:
                        print(f"处理图片失败: {e}")
                        continue
        except Exception as e:
            print(f"提取图片失败: {e}")
        
        return elements
    
    def extract_pdf_content(self, pdf_path: str) -> List[ParsedElement]:
        """从 PDF 文档提取内容"""
        doc = fitz.open(pdf_path)
        elements = []
        order = 0
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # 提取文本块（按位置排序）
            text_blocks = page.get_text("dict")["blocks"]
            page_elements = []
            
            # 处理文本
            for block in text_blocks:
                if "lines" in block:
                    block_text = ""
                    block_bbox = block["bbox"]
                    
                    for line in block["lines"]:
                        line_text = ""
                        for span in line["spans"]:
                            line_text += span["text"]
                        if line_text.strip():
                            block_text += line_text.strip() + "\n"
                    
                    if block_text.strip():
                        page_elements.append({
                            'type': 'text',
                            'content': block_text.strip(),
                            'bbox': block_bbox,
                            'page': page_num
                        })
            
            # 处理图片
            image_list = page.get_images(full=True)
            
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    # 转换为 PNG
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                    else:  # CMYK需要转换
                        pix_rgb = fitz.Pixmap(fitz.csRGB, pix)
                        img_data = pix_rgb.tobytes("png")
                        pix_rgb = None
                    
                    filename = f"pdf_img_{uuid.uuid4()}.png"
                    image_url = self.save_image(img_data, filename)
                    
                    # 获取图片位置
                    try:
                        img_rects = page.get_image_rects(img)
                        img_bbox = img_rects[0] if img_rects else [0, 0, 100, 100]
                    except:
                        img_bbox = [0, 0, 100, 100]
                    
                    page_elements.append({
                        'type': 'image',
                        'content': image_url,
                        'bbox': img_bbox,
                        'page': page_num
                    })
                    
                    pix = None
                    
                except Exception as e:
                    print(f"提取 PDF 图片失败: {e}")
                    continue
            
            # 按位置排序（从上到下，从左到右）
            page_elements.sort(key=lambda x: (x['bbox'][1], x['bbox'][0]))
            
            # 转换为 ParsedElement 对象
            for element in page_elements:
                elements.append(ParsedElement(
                    element_type=element['type'],
                    content=element['content'],
                    position=tuple(element['bbox']),
                    page=element['page'],
                    order=order
                ))
                order += 1
        
        doc.close()
        return elements
    
    def smart_split_questions(self, elements: List[ParsedElement]) -> List[ParsedQuestion]:
        """智能切分题目 - 改进的算法，专门处理试卷格式"""
        # 分离文本和图片元素
        text_elements = [e for e in elements if e.element_type == 'text']
        image_elements = [e for e in elements if e.element_type == 'image']
        
        if not text_elements:
            return []
        
        # 合并所有文本内容
        full_text = "\n".join(e.content for e in text_elements)
        
        print(f"[DEBUG] 提取的文本总长度: {len(full_text)}")
        print(f"[DEBUG] 图片数量: {len(image_elements)}")
        print(f"[DEBUG] 文本前200字符: {repr(full_text[:200])}")
        
        # 过滤掉大题标题和非题目内容
        def is_section_header(line):
            """判断是否是大题标题"""
            line = line.strip()
            # 大题标题模式：一、二、三、等
            if re.match(r'^[一二三四五六七八九十]+[、．.]\s*[^\d]', line):
                return True
            # 其他非题目标识
            if any(keyword in line for keyword in ['班级', '姓名', '学号', '共.*分', '每题.*分', '知识比拼']):
                return True
            # 空白填写行
            if re.match(r'^[_]+$', line) or '______' in line:
                return True
            return False
        
        def is_question_number(line):
            """判断是否是题目编号"""
            line = line.strip()
            # 标准题目编号：数字. 数字、 数字)
            if re.match(r'^\d{1,2}[\.．、\)\)]\s*', line):
                return True
            # 括号编号：(数字)
            if re.match(r'^[\(（]\d{1,2}[\)）]\s*', line):
                return True
            return False
        
        # 题目编号匹配模式（更精确）
        question_patterns = [
            (r'^(\d{1,2})[\.．]\s*(.*)', 'number_dot'),        # 1. 题目内容
            (r'^(\d{1,2})、\s*(.*)', 'number_comma'),          # 1、题目内容  
            (r'^[\(（](\d{1,2})[\)）]\s*(.*)', 'paren'),    # (1) 题目内容
            (r'^(\d{1,2})\)\s*(.*)', 'number_paren'),           # 1) 题目内容
        ]
        
        questions = []
        lines = [line.strip() for line in full_text.split('\n') if line.strip()]
        current_question_num = None
        current_content = []
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # 跳过大题标题和非题目内容
            if is_section_header(line):
                i += 1
                continue
                
            # 尝试匹配题目编号
            matched = False
            for pattern, pattern_name in question_patterns:
                match = re.match(pattern, line)
                if match:
                    # 保存上一道题
                    if current_question_num is not None and current_content:
                        content = '\n'.join(current_content).strip()
                        if len(content) > 5:  # 过滤太短的内容
                            questions.append(ParsedQuestion(
                                question_no=current_question_num,
                                question_text=content,
                                image_urls=[],
                                raw_content=content,
                                confidence=0.9
                            ))
                    
                    # 开始新题
                    try:
                        current_question_num = int(match.group(1))
                        current_content = [line]  # 包含题目编号的完整行
                        matched = True
                        break
                    except ValueError:
                        continue
            
            if not matched:
                # 如果不是题目编号，添加到当前题目内容
                if current_question_num is not None:
                    current_content.append(line)
                else:
                    # 如果还没开始任何题目，检查是否是没有编号的题目
                    if len(line) > 20 and not any(char in line for char in ['：', ':', '（', '(']):
                        current_question_num = len(questions) + 1
                        current_content = [line]
            
            i += 1
        
        # 处理最后一题
        if current_question_num is not None and current_content:
            content = '\n'.join(current_content).strip()
            if len(content) > 5:
                questions.append(ParsedQuestion(
                    question_no=current_question_num,
                    question_text=content,
                    image_urls=[],
                    raw_content=content,
                    confidence=0.9
                ))
        
        print(f"[DEBUG] 第一次识别到题目数量: {len(questions)}")
        
        # 如果没有识别到题目，尝试更灵活的策略
        if not questions:
            print("[DEBUG] 尝试灵活策略识别题目")
            # 按行分析，寻找可能的题目
            potential_questions = []
            for line in lines:
                if is_section_header(line):
                    continue
                # 寻找看起来像题目的行（包含问号、填空、选择等）
                if (len(line) > 15 and 
                    ('？' in line or '?' in line or 
                     '____' in line or '（）' in line or 
                     '填空' in line or '计算' in line or 
                     '选择' in line or '解答' in line)):
                    potential_questions.append(line)
            
            # 创建题目
            for i, question_text in enumerate(potential_questions[:30]):  # 最多30题
                questions.append(ParsedQuestion(
                    question_no=i + 1,
                    question_text=question_text,
                    image_urls=[],
                    raw_content=question_text,
                    confidence=0.7
                ))
        
        # 清理和重新编号
        filtered_questions = []
        for i, q in enumerate(questions):
            # 过滤太短或无意义的题目
            if len(q.question_text.strip()) > 8:
                q.question_no = len(filtered_questions) + 1
                filtered_questions.append(q)
        
        questions = filtered_questions
        
        # 智能图片分配：根据题目位置和内容分配图片
        if image_elements and questions:
            # 策略1：按题目数量平均分配，但优先给有图片需求的题目
            questions_need_images = []
            questions_no_images = []
            
            for question in questions:
                # 判断题目是否需要图片（包含“图”、“看图”等关键词）
                if any(keyword in question.question_text for keyword in ['图', '看图', '绘图', '示意图', '线段图']):
                    questions_need_images.append(question)
                else:
                    questions_no_images.append(question)
            
            # 先给有明确图片需求的题目分配
            current_img_idx = 0
            if questions_need_images:
                images_per_special = max(1, len(image_elements) // 2 // len(questions_need_images))
                for question in questions_need_images:
                    if current_img_idx < len(image_elements):
                        end_idx = min(current_img_idx + images_per_special, len(image_elements))
                        question.image_urls = [img.content for img in image_elements[current_img_idx:end_idx]]
                        current_img_idx = end_idx
            
            # 剩余图片分配给其他题目
            remaining_images = len(image_elements) - current_img_idx
            if remaining_images > 0 and questions_no_images:
                images_per_normal = max(1, remaining_images // len(questions_no_images))
                for i, question in enumerate(questions_no_images):
                    if current_img_idx < len(image_elements):
                        img_count = images_per_normal
                        # 最后一题获得所有剩余图片
                        if i == len(questions_no_images) - 1:
                            img_count = len(image_elements) - current_img_idx
                        end_idx = min(current_img_idx + img_count, len(image_elements))
                        question.image_urls = [img.content for img in image_elements[current_img_idx:end_idx]]
                        current_img_idx = end_idx
            
            # 如果没有特殊需求，就平均分配
            if not questions_need_images:
                current_img_idx = 0
                images_per_question = len(image_elements) // len(questions) if questions else 0
                extra_images = len(image_elements) % len(questions) if questions else 0
                
                for i, question in enumerate(questions):
                    img_count = images_per_question
                    if i < extra_images:
                        img_count += 1
                    
                    if img_count > 0:
                        end_idx = min(current_img_idx + img_count, len(image_elements))
                        question.image_urls = [img.content for img in image_elements[current_img_idx:end_idx]]
                        current_img_idx = end_idx
        
        print(f"[DEBUG] 最终题目数量: {len(questions)}")
        for i, q in enumerate(questions[:5]):  # 显示前5道题用于调试
            print(f"[DEBUG] 题目{i+1}: {q.question_text[:50]}...")
            
        return questions
    
    def parse_document(self, file_path: str, file_type: str, split_method: str = 'auto') -> List[ParsedQuestion]:
        """解析文档主方法 - 支持多种切分方案
        
        Args:
            file_path: 文件路径
            file_type: 文件类型 (.docx 或 .pdf)
            split_method: 切分方法
                - 'auto': 自动选择最佳方法
                - 'structure': 基于文档结构
                - 'visual': 基于视觉布局
                - 'smart': 原智能逻辑切分
                - 'manual': 手动辅助模式
        """
        try:
            print(f"[DEBUG] 开始解析文档: {file_path}, 类型: {file_type}, 方法: {split_method}")
            
            questions = []
            
            # 根据方法选择切分策略
            if split_method == 'visual':
                questions = self.visual_based_split(file_path)
                
            elif split_method == 'structure':
                if file_type == '.docx':
                    elements = self.extract_word_content_structured(file_path)
                    questions = self.structure_based_split(elements)
                elif file_type == '.pdf':
                    elements = self.extract_pdf_content(file_path)
                    questions = self.structure_based_split(elements)
                    
            elif split_method == 'smart':
                # 原有的智能逻辑切分
                if file_type == '.docx':
                    elements = self.extract_word_content_structured(file_path)
                elif file_type == '.pdf':
                    elements = self.extract_pdf_content(file_path)
                else:
                    raise ValueError(f"不支持的文件类型: {file_type}")
                questions = self.smart_split_questions(elements)
                
            elif split_method == 'auto' or split_method == 'manual':
                # 自动选择或手动模式：尝试多种方法
                methods_to_try = ['structure', 'visual', 'smart']
                best_result = []
                best_confidence = 0
                
                for method in methods_to_try:
                    try:
                        temp_questions = self.parse_document(file_path, file_type, method)
                        if temp_questions:
                            avg_confidence = sum(q.confidence for q in temp_questions) / len(temp_questions)
                            print(f"[DEBUG] 方法 {method}: {len(temp_questions)} 道题, 平均置信度: {avg_confidence:.2f}")
                            
                            # 选择最佳结果
                            if avg_confidence > best_confidence or not best_result:
                                best_result = temp_questions
                                best_confidence = avg_confidence
                    except Exception as e:
                        print(f"[DEBUG] 方法 {method} 失败: {e}")
                        continue
                
                questions = best_result
            
            else:
                raise ValueError(f"不支持的切分方法: {split_method}")
            
            print(f"[DEBUG] 解析完成，共 {len(questions)} 道题")
            return questions
            
        except Exception as e:
            print(f"解析文档失败: {e}")
            raise

# 创建解析器实例
parser = DocumentParser()

@router.post("/api/parse/document")
async def parse_document(file: UploadFile = File(...), paper_id: int = Form(...), split_method: str = Form('auto')):
    """解析Word/PDF文档并切分题目
    
    Args:
        file: 上传的文档文件
        paper_id: 试卷ID
        split_method: 切分方法 ('auto', 'structure', 'visual', 'smart', 'manual')
    """
    
    # 验证文件格式
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")
        
    file_extension = os.path.splitext(file.filename)[1].lower()
    if file_extension not in SUPPORTED_FORMATS:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式。支持格式: {', '.join(SUPPORTED_FORMATS.keys())}")
    
    # 验证切分方法
    valid_methods = ['auto', 'structure', 'visual', 'smart', 'manual']
    if split_method not in valid_methods:
        split_method = 'auto'  # 默认值
    
    # 验证文件大小
    file.file.seek(0, 2)
    file_size = file.file.tell()
    file.file.seek(0)
    
    max_size = 10 * 1024 * 1024  # 10MB
    if file_size > max_size:
        raise HTTPException(status_code=400, detail="文件大小不能超过10MB")
    
    temp_path = None
    try:
        print(f"[API] 开始处理文件: {file.filename}, 切分方法: {split_method}")
        
        # 保存临时文件
        temp_path = parser.save_temp_file(file)
        print(f"[API] 临时文件保存到: {temp_path}")
        
        # 解析文档
        questions = parser.parse_document(temp_path, file_extension, split_method)
        
        print(f"[API] 解析完成，共 {len(questions)} 道题")
        
        # 检查是否识别到题目
        if not questions:
            return {
                "success": False,
                "message": "未识别到题目",
                "suggestions": [
                    "请检查文档格式，确保题目有明确的编号（如1. 2. 3.）",
                    "支持的题目格式：1. 题目内容、第1题、(1) 题目内容、一、题目内容",
                    "或尝试手动添加题目",
                    f"当前使用的切分方法: {split_method}, 可尝试其他方法"
                ],
                "questions": [],
                "total_count": 0,
                "file_info": {
                    "filename": file.filename,
                    "size": file_size,
                    "type": SUPPORTED_FORMATS[file_extension]
                },
                "split_method": split_method
            }
        
        # 转换为API响应格式
        result_questions = []
        for q in questions:
            result_questions.append({
                "question_no": q.question_no,
                "question_text": q.question_text,
                "image_urls": q.image_urls,
                "raw_content": q.raw_content,
                "confidence": q.confidence
            })
        
        return {
            "success": True,
            "questions": result_questions,
            "total_count": len(result_questions),
            "file_info": {
                "filename": file.filename,
                "size": file_size,
                "type": SUPPORTED_FORMATS[file_extension]
            },
            "split_method": split_method,
            "avg_confidence": sum(q.confidence for q in questions) / len(questions) if questions else 0
        }
        
    except Exception as e:
        print(f"[API] 解析失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"解析失败: {str(e)}")
    
    finally:
        # 清理临时文件
        if temp_path:
            parser.cleanup_temp_file(temp_path)

@router.get("/api/parse/test")
async def test_parse_api():
    """测试解析API是否正常工作"""
    return {"message": "文档解析API正常工作", "supported_formats": SUPPORTED_FORMATS}